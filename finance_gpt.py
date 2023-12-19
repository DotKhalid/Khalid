from flask import Flask, render_template, request, send_file, Blueprint
from langchain.llms import OpenAI
import docx
import datetime
import time
import base64

finance_gpt_app = Blueprint('finance_gpt', __name__)

# Define a placeholder for the OpenAI API key
openai_api_key = None


@finance_gpt_app.route('/', methods=['GET', 'POST'])
def financial_report():
    global openai_api_key

    if request.method == 'POST':
        openai_api_key = request.form.get('openai_api_key')
        start_up_name = request.form.get('start_up_name')
        start_up_description = request.form.get('start_up_description')
        country = request.form.get('country')
        sector = request.form.getlist('sector')
        funding = request.form.getlist('funding')

        if openai_api_key.startswith('sk-'):
            date_today = datetime.date.today()
            Funding_generated_summary = generate_response(f"I'm currently in the process of exploring funding options for my startup named {start_up_name}, and I'd like to gather as much information as possible. The description of my startup is {start_up_description}. \n\
                I'm particularly interested in understanding the various funding sources available to early-stage startups like mine and any specific tips or considerations. I've selected the following {funding} options. \n\
                To begin, I'd appreciate an overview of the different types of funding sources that are accessible to my startups in {country} and related to {', '.join(sector)}. Moreover, I'd like to understand the eligibility requirements and criteria that startups typically need to meet for each funding source I've selected. This information will be invaluable as I evaluate which options align with \n\
                my startup's current stage and objectives. Preparing strong applications or pitches is crucial when seeking funding. Therefore, I would welcome any advice or tips on how to present a compelling case to potential investors or funding organizations. \n\
                Understanding what investors look for can significantly enhance my chances of securing the necessary funds. Networking is often a vital aspect of the fundraising process. \n\
                If you could provide strategies for building connections with potential investors or organizations that provide funding, I would greatly appreciate it. Insights into effective networking can be a game-changer. \n\
                Additionally, I'd like to be aware of common challenges or pitfalls that startups frequently encounter during the fundraising process. Learning from these experiences can help me avoid potential setbacks and navigate the process more effectively. \n\
                Lastly, timing and planning are critical considerations. Insights into when it's the right time to seek funding and how to plan for a successful fundraising campaign would be highly valuable. \n\
                If you could also share any relevant resources, articles, or additional advice on this topic, it would be greatly appreciated. Your assistance in this matter is of utmost importance to me as I embark on this funding journey for my startup. Put it in point form and complete each point and up-to-date specified information.")

            Legal_generated_summary = generate_response(f"I am seeking your legal expertise to guide me through the process of launching my startup called {start_up_name} in a specific {country}. I would appreciate comprehensive advice that covers all relevant legal requirements, regulations, and considerations unique to this jurisdiction and related to {', '.join(sector)}. \n\
                Please provide as much information as possible to ensure a successful and compliant startup launch in this country. Your insights are invaluable in navigating the legal landscape effectively. Put it in point form and complete each point and up-to-date specified information.")

            generate_report(start_up_name, country, date_today, Funding_text_summary=Funding_generated_summary,
                            Legal_text_summary=Legal_generated_summary)

    return render_template('financial_report.html')


def generate_report(Company_name, country, report_date, Funding_text_summary=None, Legal_text_summary=None):
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading("Financial Report", 0)
    doc.add_paragraph('Authored By: MoneyMentor FinGPT LLM')
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Created For: {Company_name}')
    doc.add_paragraph(f'Country based: {country}')
    doc.add_heading(f'Navigating the Intersection: A Comprehensive Guide to {Company_name} Financial and Legal Strategies')

    # Funding Strategies
    doc.add_heading('Funding Strategies')
    doc.add_paragraph(Funding_text_summary)

    # Legal Strategies
    doc.add_heading('Legal Strategies')
    doc.add_paragraph(Legal_text_summary)

    doc.save('Financial Report.docx')
    return send_file('Financial Report.docx', as_attachment=True)


def generate_response(input_text):
    global openai_api_key
    if openai_api_key and openai_api_key.startswith('sk-'):
        llm = OpenAI(temperature=0.3, openai_api_key=openai_api_key)
        output = llm(input_text)
        return output
    else:
        return "Please enter your OpenAI API key!"


